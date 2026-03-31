import sys
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_formatting(paragraph, text):
    """Parses basic markdown like **bold** and *italic* and applies it to the paragraph."""
    # Split text by ** for bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Check for *italic*
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    run = paragraph.add_run(sub_part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub_part)

def convert_md_to_docx(md_path):
    if not os.path.exists(md_path):
        print(f"Error: Archivo no encontrado - {md_path}")
        sys.exit(1)
        
    doc = Document()
    
    # Title styling (ETHNOS default)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()
        
    in_table = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(level=1)
            apply_formatting(p, line[2:])
        elif line.startswith('## '):
            p = doc.add_heading(level=2)
            apply_formatting(p, line[3:])
        elif line.startswith('### '):
            p = doc.add_heading(level=3)
            apply_formatting(p, line[4:])
        elif line.startswith('#### '):
            p = doc.add_heading(level=4)
            apply_formatting(p, line[5:])
        
        # Blockquotes
        elif line.startswith('>'):
            p = doc.add_paragraph()
            p.style = 'Quote'
            apply_formatting(p, line[1:].strip())
            
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            apply_formatting(p, line[2:])
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            # Remove the number for the style to handle it, or just keep it
            text = re.sub(r'^\d+\.\s', '', line)
            apply_formatting(p, text)
            
        # Tables (Very basic markdown table parsing)
        elif line.startswith('|'):
            # Just add as monospace plain text for now to avoid breaking
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            
        # Horizontal rules
        elif line == '---':
            doc.add_paragraph('_' * 50)
            
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            apply_formatting(p, line)

    # Save logic
    dir_name = os.path.dirname(md_path)
    base_name = os.path.basename(md_path)
    file_name_without_ext = os.path.splitext(base_name)[0]
    out_path = os.path.join(dir_name, f"{file_name_without_ext}.docx")
    
    try:
        doc.save(out_path)
        print(f"Éxito: Reporte convertido a DOCX y guardado en {out_path}")
    except Exception as e:
        print(f"Error al guardar DOCX: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python ethnos_docx_exporter.py <ruta_del_archivo.md>")
        sys.exit(1)
    
    md_file = sys.argv[1]
    convert_md_to_docx(md_file)
